from re import A
from django.http.response import HttpResponse
from rest_framework import viewsets, permissions, status
from rest_framework.decorators import api_view, permission_classes, authentication_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from django.contrib.auth.models import User, Group
from django.contrib.auth import authenticate, login, logout
from django.conf import settings
from django.core.mail import send_mail
from django.http import FileResponse
from django.db.models import Sum
from hotelifba.models import Empresa, Empregado, Estatistica, Quarto, Cliente, Servico, Reserva, Estadia
from hotelifba.serializer import QuartoSerializer, ServicoSerializer, EstatisticaSerializer, ClienteSerializer, ReservaSerializer, EstadiaSerializer, EmpregadoSerializer, UserSerializer, EmpresaSerializer, GroupSerializer
import string
import secrets
from datetime import datetime
from docx import Document
from drf_yasg import openapi
from drf_yasg.utils import swagger_auto_schema
from django.contrib.auth.decorators import login_required
from rest_framework.authtoken.models import Token
from rest_framework.authentication import TokenAuthentication


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all().order_by('-date_joined')
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]


class GroupViewSet(viewsets.ModelViewSet):
    queryset = Group.objects.all()
    serializer_class = GroupSerializer
    permission_classes = [permissions.IsAuthenticated]


""" View Login and Logout """


@swagger_auto_schema(methods=['POST'], request_body=UserSerializer)
@api_view(['POST'])
@permission_classes([AllowAny])
def logout_startup(request):
    try:
        logout(request)
        return HttpResponse("Logout realizado com sucesso", status=status.HTTP_200_OK)
    except KeyError:
        return HttpResponse('Não foi possível realizar o login', status=status.HTTP_404_NOT_FOUND)


user_response = openapi.Response('Response Description', UserSerializer)

@swagger_auto_schema(methods=['POST'], request_body=UserSerializer)
@api_view(['POST'])
@permission_classes([AllowAny])
def login_startup(request):

    if request.method == 'POST':
        serializer = UserSerializer(data=request.data)

        password = request.data['password']
        username = request.data['username']

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
    
            try:
                token = Token.objects.get(user_id=user.id)

            except Token.DoesNotExist:
                token = Token.objects.create(user=user)

            client = Cliente.objects.get(nome = username)

            return Response({
                'token': token.key,
                'user_id': client.pk,
                'user_username': user.username,
                'email': user.email
            }, status=status.HTTP_200_OK)
        else:
            return Response([], status=status.HTTP_401_UNAUTHORIZED)


def verificaPermissaoConfiguracao(request):
    usernameLogado = request.user.username
    funcionario = Empregado.objects.get(login=usernameLogado)
    print(f'{funcionario.funcao}')

    if funcionario.funcao == "G" or request.user.is_superuser:  # se ele for gerente ou superusuario
        return True
    return False


""" View Empresas """
empresa_response = openapi.Response('Response Description', EmpresaSerializer)


@swagger_auto_schema(method='GET', responses={200: empresa_response})
@swagger_auto_schema(methods=['POST'], request_body=EmpresaSerializer)
@api_view(['GET', 'POST'])
@permission_classes([AllowAny])
def empresa_list(request):
    if request.method == 'GET':
        empresa = Empresa.objects.all()
        serializer = EmpresaSerializer(empresa, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = EmpresaSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(methods=['PUT'], request_body=EmpresaSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def empresa_detail(request, pk):
    try:
        empresa = Empresa.objects.get(pk=pk)
    except Empresa.DoesNotExist:
        return Response('Empresa não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = EmpresaSerializer(empresa)
        return Response(serializer.data)

    elif request.method == 'PUT':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = EmpresaSerializer(empresa, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            empresa.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)


""" View Quarto """
quartos_response = openapi.Response('Response Description', QuartoSerializer)


@swagger_auto_schema(method='GET', responses={200: quartos_response})
@swagger_auto_schema(methods=['POST'], request_body=QuartoSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def quarto_list(request):
    if request.method == 'GET':
        quarto = Quarto.objects.all()
        serializer = QuartoSerializer(quarto, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = QuartoSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(methods=['PUT'], request_body=QuartoSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def quarto_detail(request, pk):
    try:
        quarto = Quarto.objects.get(pk=pk)
    except Quarto.DoesNotExist:
        return Response('Quarto não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = QuartoSerializer(quarto)
        return Response(serializer.data)

    elif request.method == 'PUT':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = QuartoSerializer(quarto, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            quarto.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)


def quarto_disponivel_capacidade_support(capacidade):
    quarto = Quarto.objects.filter(capacidade__icontains=capacidade, isDisponivel=True).first()
    return quarto.id


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def quarto_disponivel_capacidade(request, capacidade):

    if request.method == 'GET':
        quarto = Quarto.objects.filter(capacidade__icontains=capacidade, isDisponivel=True).first()
        serializer = QuartoSerializer(quarto)
            
        return Response(serializer.data)

# Altera status do quarto quando o mesmo é reservado


def update_quarto_reserva(pk):
    quarto = Quarto.objects.get(pk=pk)

    if quarto.isDisponivel is True:  # verificação dupla
        quarto.isDisponivel = False
        quarto.save()
        return True

    return False


""" View Empregado """

empregados_response = openapi.Response(
    'Response Description', EmpregadoSerializer)


@swagger_auto_schema(method='GET', responses={200: empregados_response})
@swagger_auto_schema(methods=['POST'], request_body=EmpregadoSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def empregado_list(request):
    if request.method == 'GET':
        empregado = Empregado.objects.all()
        serializer = EmpregadoSerializer(empregado, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = EmpregadoSerializer(data=request.data)
            if serializer.is_valid():
                password = request.data['password']
                login = request.data['login']
                funcao = request.data['funcao']

                user = User.objects.create_user(login, "", password)

                if funcao == "R":
                    my_group = Group.objects.get(name='Recepcionista')

                elif funcao == "G":
                    my_group = Group.objects.get(name='Gerente')

                my_group.user_set.add(user)
                user.is_admin = True
                user.is_staff = True
                user.save()

                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(methods=['PUT'], request_body=EmpregadoSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def empregado_detail(request, pk):
    try:
        empregado = Empregado.objects.get(pk=pk)
    except Empregado.DoesNotExist:
        return Response('Empregado não encontrado', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = EmpregadoSerializer(empregado)
        return Response(serializer.data)

    elif request.method == 'PUT':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = EmpregadoSerializer(empregado, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            empregado.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)


""" View Client"""


def enviarEmail(nome, senha, email):

    send_mail(
        'Hotel IFBA - INFORMATIVO',
        f'Olá, {nome} \n\n Seja Bem Vindo(a) ao Hotel IFBA ! \n Sua conta está ativa.\n Para acessar utilize a senha :{senha}',
        settings.EMAIL_HOST_USER,
        [email]
    )


cliente_response = openapi.Response('Response Description', ClienteSerializer)


@swagger_auto_schema(method='GET', responses={200: cliente_response})
@swagger_auto_schema(methods=['POST'], request_body=ClienteSerializer)
@api_view(['GET', 'POST',])
@permission_classes([AllowAny])
def cliente_list(request):
    if request.method == 'GET':
        cliente = Cliente.objects.all()
        serializer = ClienteSerializer(cliente, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        serializer = ClienteSerializer(data=request.data)

        if serializer.is_valid():
            nome = request.data['nome']
            email = request.data['email']

            senha = string.ascii_lowercase + string.digits
            senha = ''.join(secrets.choice(senha) for i in range(8))

            if email is not None:
                enviarEmail(nome, senha, email=email)

            user = User.objects.create_user(nome, email, senha)
            my_group = Group.objects.get(name='Cliente')
            my_group.user_set.add(user)
            user.is_admin = True
            user.is_staff = True
            user.save()
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



@api_view(['GET', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def client_detail(request, pk):
    try:
        client = Cliente.objects.get(pk=pk)
    except Cliente.DoesNotExist:
        return Response('Cliente não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ClienteSerializer(client)
        return Response(serializer.data)

    elif request.method == 'DELETE':
        client.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

@api_view(['PUT'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def update(request):
    try:

        client = Cliente.objects.get(pk=request.data.get('id'))

    except Cliente.DoesNotExist:
        return Response('Cliente não encontrada', status=status.HTTP_404_NOT_FOUND)


    serializer = ClienteSerializer(client, data=request.data)
    if serializer.is_valid():
        client.email = request.data.get('email')
        client.endereco = request.data.get('endereco')
        client.telefone = request.data.get('telefone')

        client.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



""" View Servico """

servicos_response = openapi.Response('Response Description', ServicoSerializer)


@swagger_auto_schema(method='GET', responses={200: servicos_response})
@swagger_auto_schema(methods=['POST'], request_body=ServicoSerializer)
@api_view(['GET', 'POST'])
@permission_classes([AllowAny])
def servico_list(request):

    if request.method == 'GET':
        servico = Servico.objects.all()
        serializer = ServicoSerializer(servico, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = ServicoSerializer(data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(methods=['PUT'], request_body=ServicoSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def servico_detail(request, pk):
    try:
        servico = Servico.objects.get(pk=pk)
    except Servico.DoesNotExist:
        return Response('Servico não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ServicoSerializer(servico)
        return Response(serializer.data)

    elif request.method == 'PUT':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            serializer = ServicoSerializer(servico, data=request.data)
            if serializer.is_valid():
                serializer.save()
                return Response(serializer.data)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        if verificaPermissaoConfiguracao(request) is False:
            return Response("Você não tem autorização para essa operação", status=status.HTTP_403_FORBIDDEN)
        else:
            servico.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)


""" View Reserva """

reservas_response = openapi.Response('Response Description', ReservaSerializer)


@swagger_auto_schema(method='GET', responses={200: reservas_response})
@swagger_auto_schema(methods=['POST'], request_body=ReservaSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def reserva_list(request):
    if request.method == 'GET':
        reserva = Reserva.objects.all()
        serializer = ReservaSerializer(reserva, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        serializer = ReservaSerializer(data=request.data)
    
        if serializer.is_valid():

            quarto_id = request.data['quarto']
            reserva_quarto = update_quarto_reserva(quarto_id)

            if reserva_quarto is True:
                serializer.save()
                return Response(serializer.data, status=status.HTTP_201_CREATED)
            else:
                return Response("Não é possível reservar este quarto.", status=status.HTTP_400_BAD_REQUEST)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def reserva_detail(request, pk):
    try:
        reserva = Reserva.objects.get(pk=pk)
    except Reserva.DoesNotExist:
        return Response('Reserva não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ReservaSerializer(reserva)
        return Response(serializer.data)

    elif request.method == 'DELETE':
        reserva.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

@swagger_auto_schema(methods=['PUT'], request_body=ReservaSerializer)
@api_view(['PUT'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def reserva_update(request):
    try:
        reserva = Reserva.objects.get(pk=request.data.get('id'))
    except Reserva.DoesNotExist:
        return Response('Reserva não encontrada', status=status.HTTP_404_NOT_FOUND)

    serializer = ReservaSerializer(reserva, data=request.data)

    if serializer.is_valid():
        pk = request.data['quarto']
        capacidade = request.data['quantidade_pessoas']

        if capacidade > reserva.quantidade_pessoas and pk == reserva.quarto.id:
            return Response("Este quarto não comporta esta quantidade de pessoas", status=status.HTTP_400_BAD_REQUEST)

        if pk != reserva.quarto.id:  # pq houve uma alteracao de quarto
            reserva_quarto = update_quarto_reserva(pk)

            if reserva_quarto is True:
                quarto = reserva.quarto
                quarto.isDisponivel = True
                quarto.save()
                serializer.save()
                return Response(serializer.data, status=status.HTTP_200_OK)

            else:
                serializer.save()
                return Response("Este quarto já possui uma reserva", status=status.HTTP_400_BAD_REQUEST)
        else:
            serializer.save()
            return Response("Reserva atualizada", status=status.HTTP_200_OK)

    return Response("Não foi possível realizar esta operação", status=status.HTTP_400_BAD_REQUEST)

""" View Estadia """


def createReserva(request):
    cliente = Cliente.objects.get(pk=request.data['cliente'])

    reserva = Reserva()
    reserva.cartao = request.data['cartao']
    reserva.data_entrada = request.data['checkin']
    reserva.data_saida = request.data['checkout']
    reserva.cliente = cliente
    reserva.quantidade_pessoas = request.data['quantidade_pessoas']

    return reserva


def createEstadia(request):
    estadia = Estadia()
    estadia.cartao = request.data['cartao']
    estadia.checkin = request.data['checkin']
    estadia.checkout = request.data['checkout']
    estadia.data_saida_prevista = request.data['data_saida_prevista']
    estadia.quantidade_pessoas = request.data['quantidade_pessoas']
    estadia.quantidade_quartos = request.data['quantidade_quartos']
    estadia.servico = request.data['servico']

    return estadia


def updateEstadia(reservado, reserva, estadia, quarto_id):
    if reservado is True:
        quarto = Quarto.objects.get(pk=quarto_id)
        reserva.quarto = quarto
        reserva.save()

        estadia.cliente = reserva.cliente
        estadia.reserva = reserva
        estadia.save()

        estadia = Estadia.objects.get(pk=estadia.id)
        return estadia


estadias_response = openapi.Response('Response Description', EstadiaSerializer)


@swagger_auto_schema(method='GET', responses={200: estadias_response})
@swagger_auto_schema(methods=['POST'], request_body=EstadiaSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def estadia_list(request):

    if request.method == 'GET':
        estadia = Estadia.objects.all()
        serializer = EstadiaSerializer(estadia, many=True)
        return Response(serializer.data)

    elif request.method == 'POST':
        serializer = EstadiaSerializer(data=request.data)
        reserva_id = request.data['reserva']

        if reserva_id is None:
            estadia = createEstadia(request)
            reserva = createReserva(request)
            quarto_id = quarto_disponivel_capacidade_support(
                request.data['quantidade_pessoas'])

            if quarto_id is not None:

                reservado = update_quarto_reserva(quarto_id)
                estadia = updateEstadia(reservado, reserva, estadia, quarto_id)

                estadia = Estadia.objects.get(pk=estadia.id)
                serializer = EstadiaSerializer(estadia)
                return Response(serializer.data)
        else:
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(methods=['PUT'], request_body=EstadiaSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def estadia_detail(request, pk):
    try:
        estadia = Estadia.objects.get(pk=pk)
    except Estadia.DoesNotExist:
        return Response('Estadia não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = EstadiaSerializer(estadia)
        return Response(serializer.data)

    elif request.method == 'PUT':
        serializer = EstadiaSerializer(estadia, data=request.data)
        if serializer.is_valid():

            if estadia.isMudancaDeQuarto is True:
                quarto = change_room(estadia, pk)

                if quarto is not None:
                    estadia.isMudancaDeQuarto = False
                    estadia.save()
                    return Response(serializer.data)
                else:
                    return Response("Não há quartos que satisfaçam a solicitação", status=status.HTTP_404_NOT_FOUND)

            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        estadia.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)


" View checkout "


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def checkout(request, pk):
    try:
        estadia = Estadia.objects.get(pk=pk)
    except Estadia.DoesNotExist:
        return Response('Estadia não encontrada', status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        reserva = Estadia.objects.get(pk=estadia.reserva.id)

        quantidade_dias = diff_datas(estadia.checkin, estadia.checkout)
        valor_pagar = reserva.servico.preco * quantidade_dias

        document = generate_fatura(valor_pagar, quantidade_dias, pk)

        quarto = estadia.reserva.quarto
        quarto.isDisponivel = True
        quarto.save()

        if document is not None:
            createEstatistica(valor_pagar, estadia.checkin, estadia.cliente)
            return download_fatura(request, estadia.id)
        else:
            return HttpResponse(status.HTTP_404_NOT_FOUND)


def download_fatura(request, pk):
    if request.method == 'GET':
        estadia = Estadia.objects.get(pk=pk)

        nomeCliente = estadia.cliente.nome
        nomeDocumento = f'fatura{nomeCliente}.docx'
        filePath = f'hotelifba/static/media/{nomeDocumento}'

        file = open(filePath, 'rb')
        response = FileResponse(file)
        response['Content-Type'] = 'application/application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response['Content-Disposition'] = f'attachment;filename={nomeDocumento}'

        return response


def generate_fatura(valor_pagar, quantidade_dias, pk):
    estadia = Estadia.objects.get(pk=pk)

    document = Document()
    docx_title = f"fatura{estadia.cliente.nome}.docx"

    document.add_heading('Fatura', 0)

    p = document.add_paragraph('Fatura referente a estadia no Hotel IFBA ')

    document.add_heading('Informação de Estadia', 1)
    document.add_paragraph(f'Check in: {estadia.checkin}')
    document.add_paragraph(f'Check out: {estadia.checkout}')
    document.add_paragraph(f'Quantidade de dias: {quantidade_dias } dias')
    document.add_paragraph(f'Serviço prestado: {estadia.servico.tipo } ')
    document.add_paragraph(
        f'Valor diário do serviço: {estadia.servico.preco } ')
    document.add_paragraph(
        f'Quantidade de Pessoas: {estadia.quantidade_pessoas } ')

    document.add_heading('Informação do Cliente', 1)
    document.add_paragraph(f'Nome do cliente: {estadia.cliente.nome}')
    document.add_paragraph(f'Número do cartão: {estadia.cartao}')
    document.add_paragraph(f'Telefone: {estadia.cliente.telefone}')

    document.add_heading('Valor a pagar', 1)
    document.add_paragraph(f'R$: {valor_pagar}')

    document.save(f'hotelifba/static/media/{docx_title}')
    return document


def diff_datas(dataStart, dataEnd):
    dataStart = datetime.strptime(dataStart.strftime('%m/%d/%Y'), "%m/%d/%Y")
    dataEnd = datetime.strptime(dataEnd.strftime('%m/%d/%Y'), "%m/%d/%Y")

    quantidadeDias = ((dataEnd-dataStart).days)

    if quantidadeDias == 0:
        quantidadeDias = 1

    return quantidadeDias


def change_room(estadia, pk):

    quarto = Quarto.objects.filter(
        capacidade__icontains=estadia.quantidade_pessoas, isDisponivel=True).first()
    quarto_atual = estadia.reserva.quarto

    if quarto is not None:
        reserva_quarto = update_quarto_reserva(quarto.id)

        if reserva_quarto is True:
            quarto_atual.isDisponivel = True
            quarto_atual.save()
            estadia.reserva.quarto = quarto
            estadia.reserva.save()

    return quarto


""" View Estatistica """


def faturamentoAnual(valor_pagar):
    faturamentoAnual = Estatistica.objects.all().aggregate(
        Sum('faturamentoAnual'))['faturamentoAnual__sum']

    if faturamentoAnual is None:
        return valor_pagar
    else:
        return faturamentoAnual + float(valor_pagar)


def faturamentoDoTrimestre(valor_pagar, trimestreReferente):
    faturamentoDoTrimestre = Estatistica.objects.filter(trimestre=trimestreReferente).aggregate(
        Sum('faturamentoDoTrimestre'))['faturamentoDoTrimestre__sum']

    if faturamentoDoTrimestre is None:
        return valor_pagar
    else:
        return faturamentoDoTrimestre + float(valor_pagar)


def taxaOcupacaoQuartos():
    capacidadeMaxima = Quarto.objects.all().aggregate(
        Sum('capacidade'))['capacidade__sum']
    alojamentosOcupados = Quarto.objects.filter(
        isDisponivel=False).aggregate(Sum('capacidade'))['capacidade__sum']

    if alojamentosOcupados is not None:
        alojamentosOcupados = alojamentosOcupados * 100
    else:
        alojamentosOcupados = 0

    return alojamentosOcupados/capacidadeMaxima


def taxaQuartosVendidos():
    quartoVendidos = Quarto.objects.filter(isDisponivel=False).count() * 100
    quartoTotal = Quarto.objects.all().count()

    return quartoVendidos/quartoTotal


def createEstatistica(valor_pagar, checkin, cliente):
    datacheckin = datetime.strptime(checkin.strftime('%m/%d/%Y'), "%m/%d/%Y")
    anoReferente = datacheckin.year
    mes = datacheckin.month
    trimestreReferente = verificarTrimestre(mes)

    estatistica = Estatistica.objects.filter(
        trimestre=trimestreReferente, ano=anoReferente)

    if estatistica.exists():
        estatistica = Estatistica.objects.get(
            trimestre=trimestreReferente, ano=anoReferente)
        estatistica.ano = anoReferente
        estatistica.trimestre = trimestreReferente
        estatistica.faturamentoAnual = faturamentoAnual(valor_pagar)
        estatistica.faturamentoDoTrimestre = faturamentoDoTrimestre(
            valor_pagar, trimestreReferente)
        estatistica.taxaOcupacaoQuartos = taxaOcupacaoQuartos()
        estatistica.taxaQuartosVendidos = taxaQuartosVendidos()
        estatistica.custoTotalCliente = obterCustoTotalCliente(
            cliente, valor_pagar)
        estatistica.clientePremium = obterCliente(cliente)
        estatistica.clienteId = cliente.id
        estatistica.save()
    else:
        estatistica = Estatistica()
        estatistica.ano = anoReferente
        estatistica.trimestre = trimestreReferente
        estatistica.faturamentoAnual = faturamentoAnual(valor_pagar)
        estatistica.faturamentoDoTrimestre = faturamentoDoTrimestre(
            valor_pagar, trimestreReferente)
        estatistica.taxaOcupacaoQuartos = taxaOcupacaoQuartos()
        estatistica.taxaQuartosVendidos = taxaQuartosVendidos()
        estatistica.clienteId = cliente.id
        estatistica.custoTotalCliente = valor_pagar
        estatistica.clientePremium = obterCliente(cliente)
        estatistica.save()


def obterCliente(cliente):
    estatisticaDoCliente = Estatistica.objects.all().order_by(
        'custoTotalCliente').first()

    if estatisticaDoCliente is None:
        return cliente.nome
    else:
        cliente = Cliente.objects.get(pk=estatisticaDoCliente.id)
        return cliente.nome


def obterCustoTotalCliente(cliente, valor_pagar):
    estatisticaDoCliente = Estatistica.objects.filter(
        clienteId=cliente.id).order_by('custoTotalCliente').first()

    if estatisticaDoCliente is None:
        return valor_pagar
    else:
        estatisticaDoCliente.custoTotalCliente = float(
            valor_pagar) + estatisticaDoCliente.custoTotalCliente
        return estatisticaDoCliente.custoTotalCliente


def verificarTrimestre(mes):
    if mes >= 1 and mes <= 3:
        return 'primeiroTrimestre'
    if mes >= 4 and mes <= 6:
        return 'segundoTrimestre'
    if mes >= 7 and mes <= 9:
        return 'terceiroTrimestre'
    if mes >= 10 and mes <= 12:
        return 'quartoTrimestre'


estatistica_response = openapi.Response(
    'Response Description', EstatisticaSerializer)


@swagger_auto_schema(method='GET', responses={200: estatistica_response})
@api_view(['GET'])
@permission_classes([IsAuthenticated])
@authentication_classes([TokenAuthentication])
def estatistica(request):
    if request.method == 'GET':
        estatistica = Estatistica.objects.all()
        serializer = EstatisticaSerializer(estatistica, many=True)
        return Response(serializer.data)

    return Response(status=status.HTTP_400_BAD_REQUEST)
